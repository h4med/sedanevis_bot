# utils.py

import logging
import html
from functools import wraps
import io
import asyncio
import jdatetime
import pytz

from markdown_it import MarkdownIt
from bs4 import BeautifulSoup, NavigableString 
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup
)

from telegram.ext import ContextTypes
from telegram.constants import ParseMode
from telegram.error import BadRequest
from telethon import TelegramClient, errors as telethon_errors

import config
from database import SessionLocal, User, ActivityLog
from texts import Texts
from ai_services import (
    count_text_tokens
)

    
def convert_md_to_html(md_text: str, user_lang: str) -> str:
    """
    Converts Markdown to Telegram-compatible HTML with robust sanitization
    and forces text direction for each line based on the user's language.
    """
    md = MarkdownIt().disable('backticks')
    html_output = md.render(md_text)

    soup = BeautifulSoup(html_output, 'lxml')
    ALLOWED_TAGS = ['b', 'i', 's', 'u', 'code', 'pre', 'a', 'tg-spoiler']

    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        tag.name = 'b'
        tag.insert_before(NavigableString('\n'))
        tag.insert_after(NavigableString('\n'))
    for tag in soup.find_all('li'):
        tag.insert(0, NavigableString('• '))
        tag.insert_after(NavigableString('\n'))
    for tag in soup.find_all('hr'):
        tag.replace_with(NavigableString('\n---\n'))
    for tag in soup.find_all('strong'):
        tag.name = 'b'
    for tag in soup.find_all('em'):
        tag.name = 'i'

    for tag in soup.find_all(True):
        if tag.name not in ALLOWED_TAGS:
            tag.unwrap()
    if soup.body:
        final_html = soup.body.decode_contents()
    else:
        final_html = str(soup)

    final_html = re.sub(r'\n{3,}', '\n\n', final_html).strip()

    if user_lang == 'fa':
        direction_char = '\u200F' 
        lines = final_html.split('\n')
        forced_rtl_lines = [f"{direction_char}{line}" for line in lines]
        return '\n'.join(forced_rtl_lines)
    else:
        return final_html


def create_word_document(html_content: str, user_lang: str) -> io.BytesIO:
    """
    Converts HTML-formatted text to a correctly styled in-memory .docx file,
    preserving bold/italic formatting, and handling RTL direction, alignment,
    and paragraph spacing correctly.
    """
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Tahoma' if user_lang == 'fa' else 'Calibri'
    font.size = Pt(11)

    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.line_spacing = 1.15

    if user_lang == 'fa':
        p_format.right_to_left = True

    lines = html_content.split('\n')
    
    for line in lines:
        p = document.add_paragraph()

        # Set specific alignment for this paragraph
        if user_lang == 'fa':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if not line.strip():
            continue
        soup = BeautifulSoup(line, 'html.parser')
        
        for element in soup.contents:
            if isinstance(element, NavigableString):
                text = str(element)
            else:
                text = element.get_text()

            run = p.add_run(text)
            
            if user_lang == 'fa':
                run.font.rtl = True

            if not isinstance(element, NavigableString):
                if element.name in ['b', 'strong']:
                    run.bold = True
                elif element.name in ['i', 'em']:
                    run.italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer


def is_rtl_language(lang_code: str) -> bool:
    """Check if language should be displayed RTL."""
    rtl_languages = ['fa', 'ar', 'he', 'ur']
    return lang_code in rtl_languages

async def deliver_transcription_result(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    transcript_text: str,
    source_info: dict,
):
    """
    Delivers the transcription result to the user, intelligently choosing
    between sending a direct message or a file based on length.
    """
    reply_target = update.effective_message
    TELEGRAM_MESSAGE_LIMIT = 4096
    TELEGRAM_CAPTION_LIMIT = 1024
    
    context.user_data['last_text'] = transcript_text
  
    loop = asyncio.get_event_loop()
    transcription_tokens = await loop.run_in_executor(
        config.TOKEN_COUNTING_EXECUTOR,
        count_text_tokens,
        transcript_text,
        "gemini-2.0-flash-lite"
    )      
    
    action1_estimated_minutes = (transcription_tokens + 500) / config.TEXT_TOKENS_TO_MINUTES_COEFF
    action2_estimated_minutes = (transcription_tokens + 2000) / config.TEXT_TOKENS_TO_MINUTES_COEFF
    action3_estimated_minutes = (transcription_tokens + 1000) / config.TEXT_TOKENS_TO_MINUTES_COEFF
    cost_minutes_estimated = transcription_tokens / config.TEXT_TOKENS_TO_MINUTES_COEFF
    logging.info(f"Transcription tokens: {transcription_tokens}, in minutes: {cost_minutes_estimated}")
  
    # Get the user's current credit
    db_user = context.user_data.get('db_user')
    if not db_user:
        db = SessionLocal()
        db_user = db.query(User).filter(User.user_id == update.effective_user.id).first()
        db.close()
    remaining_credit = db_user.credit_minutes if db_user else 0.0
  
    # Extract information from source_info
    lang_code = source_info.get('language', 'fa')
    cost_minutes = source_info.get('cost', 0.0)

    # Determine if RTL is needed
    is_rtl = is_rtl_language(lang_code)
    context.user_data['is_rtl'] = is_rtl
    
    # Prepare information section
    if is_rtl:
        info_section = (
            f"<b>مدت:</b> {cost_minutes:.1f} دقیقه\n"
        )
        if source_info['type'] == 'media':
            info_section += f"<b>اعتبار مصرف شده:</b> {cost_minutes:.2f} دقیقه\n"
            info_section += f"<b>اعتبار باقی‌مانده:</b> {remaining_credit:.2f} دقیقه\n"
        elif source_info['type'] == 'youtube':
            info_section += f"<b>اعتبار فعلی:</b> {remaining_credit:.2f} دقیقه\n"
    else:
        info_section = (
            f"<b>Duration:</b> {cost_minutes:.1f} minutes\n"
        )
        if source_info['type'] == 'media':
            info_section += f"<b>Consumed:</b> {cost_minutes:.2f} minutes\n"
            info_section += f"<b>Remaining Credit:</b> {remaining_credit:.2f} minutes\n"
        elif source_info['type'] == 'youtube':
            info_section += f"<b>Current Credit:</b> {remaining_credit:.2f} minutes\n"
  
    # Decide delivery method based on length
    if len(transcript_text) < (TELEGRAM_MESSAGE_LIMIT - 500):
        # Short transcript -> Send as a single message
        logging.info("Transcript is short, sending as a direct message.")
        escaped_text = html.escape(transcript_text)
        
        if is_rtl:
            message_body = (
                f"{info_section}\n"
                f"✅ <b>رونوشت کامل:</b>\n\n<code>{escaped_text}</code>\n\n"
                "----------------\n"
                "👈 چه کاری روی این متن انجام دهم؟"
            )
        else:
            message_body = (
                f"{info_section}\n"
                f"✅ <b>Full Transcript:</b>\n\n<code>{escaped_text}</code>\n\n"
                "----------------\n"
                "👉 What action should I perform on this text?"
            )

        await reply_target.reply_text(
            text=message_body,
            parse_mode=ParseMode.HTML,
            reply_markup=get_action_keyboard(
                action1_estimated_minutes,
                action2_estimated_minutes,
                action3_estimated_minutes
            )
        )
    else:
        # Long transcript -> Send as a file with a preview
        logging.info("Transcript is long, sending as a file.")
        preview_text = html.escape(transcript_text[:600])
  
        if is_rtl:
            caption = (
                f"{info_section}\n"
                "🔗 فایل کامل رونوشت پیوست شد.👆\n\n"
                f"<b>پیش‌نمایش:</b>\n<code>{preview_text}</code>\n\n"
                "----------------\n"
                "👈 چه کاری روی این متن انجام دهم؟"                
            )
        else:
            caption = (
                f"{info_section}\n"
                f"<b>Preview:</b>\n<code>{preview_text}</code>\n\n"
                "----------------\n"
                "🔗 Full document attached."
            )
            
        # Safety check for caption length
        if len(caption) > TELEGRAM_CAPTION_LIMIT:
            truncate_at = TELEGRAM_CAPTION_LIMIT - 50
            if is_rtl:
                caption = caption[:truncate_at] + "... (ادامه در فایل)"
            else:
                caption = caption[:truncate_at] + "... (continued in file)"
  
        try:
            document_buffer = create_word_document(transcript_text, lang_code)
            tehran_tz = pytz.timezone('Asia/Tehran')
            current_report_time = jdatetime.datetime.now(tehran_tz).strftime("%Y%m%d-%H%M%S")
            transcription_filename = f"Transcription_{current_report_time}.docx"

            await reply_target.reply_document(
                document=document_buffer,
                filename=transcription_filename,
                caption=caption,
                parse_mode=ParseMode.HTML,
                reply_markup=get_action_keyboard(
                    action1_estimated_minutes,
                    action2_estimated_minutes,
                    action3_estimated_minutes
                ),
                read_timeout=120,
                write_timeout=120
            )
        except Exception as e:
            logging.error(f"Failed to create or send Word file: {e}", exc_info=True)
            await reply_target.reply_text(Texts.Errors.GENERIC_UNEXPECTED_ADMIN.format(error=e))


def get_or_create_user(session, user_id: int, first_name: str, username: str | None) -> tuple[User, bool]:
    """
    Retrieves a user from the database or creates a new one if they don't exist.
    Returns the user object and a boolean indicating if the user was newly created.
    """
    user = session.query(User).filter(User.user_id == user_id).first()
    if user:
        return user, False  # User exists

    # User does not exist, create a new one
    new_user = User(
        user_id=user_id,
        first_name=first_name,
        username=username,
        status='pending'  # Default status for new users
    )
    session.add(new_user)
    session.commit()
    logging.info(f"New user created in DB: {first_name} ({user_id}) with status 'pending'.")
    return new_user, True # User was newly created

def log_activity(db: SessionLocal, user_id: int, action: str, credit_change: float, details: str | None = None):
    """Creates an activity log entry for a user."""
    try:
        log_entry = ActivityLog(
            user_id=user_id,
            action_type=action,
            credit_change=credit_change,
            details=details
        )
        db.add(log_entry)
        db.commit()
        logging.info(f"Logged activity for user {user_id}: {action}, change: {credit_change}")
    except Exception as e:
        db.rollback()
        logging.error(f"Failed to log activity for user {user_id}: {e}", exc_info=True)


def check_user_status(func):
    """
    A decorator that checks user status before executing a handler.
    It handles new user creation, pending/rejected status, and credit checks.
    """
    @wraps(func)
    async def wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE, *args, **kwargs):
        effective_user = update.effective_user
        if not effective_user:
            return

        db = SessionLocal()
        try:
            user, is_new = get_or_create_user(
                session=db,
                user_id=effective_user.id,
                first_name=effective_user.first_name,
                username=effective_user.username
            )

            # If user is new, notify admin
            if is_new:
                user_details = Texts.Admin.NEW_USER_NOTIFICATION.format(
                    first_name=html.escape(user.first_name),
                    username=user.username if user.username else 'N/A',
                    user_id=user.user_id,
                    lang_code=effective_user.language_code or 'N/A'
                )
                keyboard = InlineKeyboardMarkup([
                    [
                        InlineKeyboardButton(Texts.Keyboard.APPROVE_USER, callback_data=f"approve:{user.user_id}"),
                        InlineKeyboardButton(Texts.Keyboard.REJECT_USER, callback_data=f"reject:{user.user_id}")
                    ]
                ])
                await context.bot.send_message(
                    chat_id=config.ADMIN_USER_ID,
                    text=user_details,
                    reply_markup=keyboard,
                    parse_mode=ParseMode.HTML
                )
                await update.message.reply_text(Texts.User.NEW_USER_GREETING)
                return

            # Check user status
            if user.status == 'pending':
                await update.message.reply_text(Texts.User.PENDING_STATUS)
                return
            if user.status in ['rejected', 'banned']:
                await update.message.reply_text(Texts.User.REJECTED_STATUS)
                return
            
            # If user is approved, add user object to context and proceed
            if user.status == 'approved':
                context.user_data['db_user'] = user
                return await func(update, context, *args, **kwargs)

        finally:
            db.close()

    return wrapper


# def get_action_keyboard():
def get_action_keyboard(action1_estimated_minutes=0, action2_estimated_minutes=0, action3_estimated_minutes=0):

    def format_minutes(minutes):
        if minutes >= 1:
            return f"{minutes:.1f}" if minutes != int(minutes) else f"{int(minutes)}"
        else:
            return f"{minutes:.2f}"

    keyboard = [
        [
            InlineKeyboardButton(
                f"{Texts.Keyboard.SUMMARY_SHORT} (هزینه ~ {format_minutes(action1_estimated_minutes)}m)", 
                callback_data='summary_short'
            )
        ],
        [
            InlineKeyboardButton(
                f"{Texts.Keyboard.EXTRACT_POINTS} (هزینه ~ {format_minutes(action2_estimated_minutes)}m)", 
                callback_data='extract_points'
            )
        ],
        [
            InlineKeyboardButton(
                f"{Texts.Keyboard.EXTRACT_MINUTES} (هزینه ~ {format_minutes(action3_estimated_minutes)}m)", 
                callback_data='extract_mom'
            )
        ]        
    ]    
    return InlineKeyboardMarkup(keyboard)

# --- Telethon Helper ---
async def ensure_telethon_client():
    if config.telethon_client is None or not config.telethon_client.is_connected():
        logging.info("Initializing Telethon client...")

        # --- Point the session file to the persistent data directory ---
        session_path = 'persistent_data/bot_session_name'        

        # client = TelegramClient('bot_session_name', config.TG_API_ID, config.TG_API_HASH, auto_reconnect=True)
        client = TelegramClient(session_path, config.TG_API_ID, config.TG_API_HASH, auto_reconnect=True)
        await client.start(bot_token=config.TG_BOT_TOKEN)
        config.telethon_client = client
        logging.info("Telethon client started and authorized.")
    return config.telethon_client

async def extract_text_from_docx(update: Update, context: ContextTypes.DEFAULT_TYPE) -> str:
    """Extract text from DOCX file."""
    try:
        # Get file info
        docx_file = update.message.document
        
        # Check file size (optional - prevent very large files)
        if docx_file.file_size > 5 * 1024 * 1024:  # 5MB limit
            await update.message.reply_text("File too large. Please upload a file smaller than 5MB.")
            return None
        
        # Get file object
        file_obj = await context.bot.get_file(docx_file.file_id)
        
        # Download file content
        file_bytes = await file_obj.download_as_bytearray()
        
        # Extract text from DOCX
        doc = Document(io.BytesIO(bytes(file_bytes)))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
        
    except Exception as e:
        await update.message.reply_text("Sorry, I couldn't process your DOCX file. Please make sure it's a valid Word document.")
        print(f"Error processing DOCX: {e}")
        return None